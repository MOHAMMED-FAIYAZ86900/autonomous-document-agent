"""
Microsoft Word document generator.
"""

from pathlib import Path

from docx import Document

from app.agent.state import AgentState
from app.core.config import settings
from app.core.logging import get_logger


class DocumentGenerator:
    """
    Generates Microsoft Word documents from
    AI-generated content.
    """

    def __init__(self) -> None:
        self.logger = get_logger(__name__)

    def generate(
        self,
        state: AgentState,
    ) -> AgentState:
        """
        Generate a DOCX document from the
        generated content.
        """

        self.logger.info(
            "Document generation started."
        )

        try:

            document = Document()

            document.add_heading(
                "AI Generated Document",
                level=1,
            )

            document.add_paragraph(
                state.generated_content or ""
            )

            output_dir = (
                settings.STORAGE_DIR
                / "generated_docs"
            )

            output_dir.mkdir(
                parents=True,
                exist_ok=True,
            )

            filename = "generated_document.docx"

            from app.storage.manager import StorageManager

            storage = StorageManager()

            output_path = storage.save_document(
                document=document,
                filename=filename,
            )

            state.output_file = str(output_path)

            self.logger.info(
                "Document saved successfully."
            )

            return state

        except Exception as exc:

            self.logger.exception(
                "Document generation failed."
            )

            state.status = "failed"

            state.errors.append(
                str(exc)
            )

            return state