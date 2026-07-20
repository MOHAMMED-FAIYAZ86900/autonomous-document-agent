"""
Document Builder
"""

from docx import Document
from docx.shared import Pt


class DocumentBuilder:
    """Builds Microsoft Word documents."""

    def __init__(self):
        self.document = Document()

    def add_title(self, title: str):
        heading = self.document.add_heading(level=1)
        run = heading.add_run(title)
        run.font.size = Pt(20)

    def add_content(self, content: str):
        self.document.add_paragraph(content)

    def build(self) -> Document:
        return self.document