"""
Unit tests for StorageManager.
"""

from docx import Document

from app.storage.manager import StorageManager


def test_save_document(tmp_path):
    """
    Test saving a DOCX document.
    """

    storage = StorageManager(storage_path=tmp_path)

    document = Document()
    document.add_heading("Hello World")

    filename = "test.docx"

    output = storage.save_document(
        document=document,
        filename=filename,
    )

    assert output.exists()
    assert output.name == filename


def test_exists(tmp_path):
    """
    Test exists().
    """

    storage = StorageManager(storage_path=tmp_path)

    document = Document()

    filename = "exists.docx"

    storage.save_document(
        document,
        filename,
    )

    assert storage.exists(filename)


def test_get_path(tmp_path):
    """
    Test get_path().
    """

    storage = StorageManager(storage_path=tmp_path)

    filename = "sample.docx"

    path = storage.get_path(filename)

    assert path.name == filename
    assert path.parent == tmp_path


def test_list_documents(tmp_path):
    """
    Test list_documents().
    """

    storage = StorageManager(storage_path=tmp_path)

    for i in range(3):

        document = Document()

        storage.save_document(
            document,
            f"doc_{i}.docx",
        )

    documents = storage.list_documents()

    assert len(documents) == 3
